from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
from docxtpl import DocxTemplate
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import tempfile, os, uuid, re, json
from pathlib import Path
from fastapi.staticfiles import StaticFiles

load_dotenv()

app = FastAPI()

# Allow CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Replace with frontend domain in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = tempfile.gettempdir()
PLACEHOLDER_REGEX = re.compile(r"\[([^\]]+)\]|\$\[_{5,}\]")

# -------------------
# LangChain setup
# -------------------
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)

# Memory: track conversation history safely with a single input key
memory = ConversationBufferMemory(
    input_key="legal_context",  
    memory_key="history",
    return_messages=True
)

# Single input key prompt template
prompt = PromptTemplate(
    input_variables=["legal_context"],
    template="{legal_context}"
)

chain = LLMChain(llm=llm, prompt=prompt, memory=memory)

# -------------------
# Prompt registry loader
# -------------------
PROMPTS_DIR = "prompts"

def load_prompt(doc_type: str) -> str:
    path = os.path.join(PROMPTS_DIR, f"{doc_type}.txt")
    if not os.path.exists(path):
        return "You are a helpful assistant helping fill out a legal document."
    with open(path, "r") as f:
        return f.read()

# -------------------
# Helper function: convert brackets to {{ placeholder }}
# -------------------
def convert_to_docxtpl_format(text: str) -> str:
    """
    Convert placeholders like [Investor Name] or $[_____]
    into docxtpl Jinja-style placeholders {{ Investor_Name }}
    """
    def replacer(match):
        raw = match.group(1).strip()
        key = raw.replace(" ", "_")  # replace spaces with underscores
        return f"{{{{ {key} }}}}"
    
    return re.sub(r'\$?\[([^\]]+)\]', replacer, text)

# -------------------
# Upload document & extract placeholders
# -------------------
@app.post("/upload")
async def upload_doc(file: UploadFile = File(...), doc_type: str = Form("generic")):
    file_id = str(uuid.uuid4())
    tmp_path = os.path.join(UPLOAD_DIR, file_id + ".docx")

    with open(tmp_path, "wb") as f:
        f.write(await file.read())

    doc = Document(tmp_path)

    # Convert placeholders to {{ ... }} format
    for para in doc.paragraphs:
        para.text = convert_to_docxtpl_format(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = convert_to_docxtpl_format(cell.text)

    # Save converted doc
    doc.save(tmp_path)

    # Extract placeholders for frontend
    PLACEHOLDER_REGEX = re.compile(r"\{\{\s*([^\}]+)\s*\}\}")
    placeholders = set()
    for para in doc.paragraphs:
        for m in PLACEHOLDER_REGEX.finditer(para.text):
            placeholders.add(m.group(1))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in PLACEHOLDER_REGEX.finditer(cell.text):
                    placeholders.add(m.group(1))

    return {
        "file_id": file_id,
        "placeholders": list(placeholders),
        "doc_type": doc_type
    }

# -------------------
# Chat endpoint - Smart version
# -------------------
@app.post("/chat")
async def chat_flow(data: dict = Body(...)):
    placeholders = data.get("placeholders", [])
    answers = data.get("answers", {})
    doc_type = data.get("doc_type", "generic")

    # Find the next unfilled placeholder
    unfilled_phs = [ph for ph in placeholders if ph not in answers]
    if not unfilled_phs:
        return {"done": True, "message": "✅ All placeholders filled!"}

    next_ph = unfilled_phs[0]  # pick the first unfilled

    # Load legal-specific context
    legal_context_text = load_prompt(doc_type)

    # Combine all relevant info into a single input string
    combined_input = (
        f"{legal_context_text}\n\n"
        f"Conversation so far:\n{memory.buffer_as_str}\n\n"
        f"Next placeholder to fill: {next_ph}\n"
        f"Already filled placeholders:\n{json.dumps(answers, indent=2)}\n\n"
        f"Ask the user a clear, natural question to fill this placeholder."
    )

    # Run LangChain to generate the next question
    question = chain.run(legal_context=combined_input)

    return {
        "done": False,
        "placeholder": next_ph,
        "message": question
    }

# -------------------
# Fill placeholders
# -------------------
@app.post("/fill")
async def fill_doc(file_id: str = Form(...), values: str = Form(...)):
    values_dict = json.loads(values)
    src_path = os.path.join(UPLOAD_DIR, file_id + ".docx")

    if not os.path.exists(src_path):
        return JSONResponse({"error": "File not found"}, status_code=404)

    tpl = DocxTemplate(src_path)
    tpl.render(values_dict)  # values_dict keys must match {{ placeholders }}
    out_path = os.path.join(UPLOAD_DIR, file_id + "_filled.docx")
    tpl.save(out_path)

    return FileResponse(out_path, filename="filled.docx")

# -------------------
# Serve React frontend
# -------------------
build_path = Path(__file__).parent.parent / "frontend" / "build"
app.mount("/static", StaticFiles(directory=build_path / "static"), name="static")

@app.get("/{full_path:path}")
async def serve_react(full_path: str):
    index_file = build_path / "index.html"
    return FileResponse(index_file)
